'''
Automatically locates files within the 'messages' folder. Writes .docx instead of .txt files if python-docx is available. 

Python
Ability to read .docx files; i.e. LibreOffice/Microsoft Word. Google Docs can't read .docx files of this size.
The .properties files to be translated. 

This program takes two .properties translation files:
    A, the complete version
    B, an incomplete version that needs to be filled in

It does the following:
    Fills in missing entries in file B by copying them from file A (so B has every key A does).
    Preserves any lines in B that aren’t key-value pairs (like comments or extra formatting).
    Reorders all entries in B to match the exact key order of A — so both files are perfectly aligned, line-by-line.

Output is written to a new file in the same folder as this program.

After you receive this output, you are expected to convert the .txt file to a .docx file, and translate it to your
desired language with Google Translate and then review manually afterwards if desired. Overwrite the text in your
original .properties file with that of the .docx file. 
'''

# It turns out that... ugh. There's another error. 
import os
import sys
import time
has_docx = True
#try:
from docx import Document
#except ImportError:
#    print("You don't have python-docx. (pip install python-docx) The output will be a '.txt'.")
#    has_docx = False
suffix = ".docx" if has_docx else ".txt"
path = "C:\\Users\\awang\\Documents\\aiden\\redundant\\V0.8.4.9\\messages" # go into all subfolders with depth=1 and per

def make_dictionary(filepath):
    file = open(filepath, "r", encoding = "utf-8") 
    text = file.readlines() 
    file.close(); 
    # keys = {}
    keys = []
    values = []
    for i in text:
        index = i.find("=")
        # attempt to add line-by-line translation, maintain structure accurate to chinese (default) translation
        if index == -1:
            keys.append(i)
            values.append("")
            continue
        # keys[i[:index]] = i[index+1:]
        keys.append(i[:index])
        values.append(i[index+1:-1]) # strip newline hack??
    return (keys, values)

def perform(path1, path2, path3):
    keys1, values1 = make_dictionary(path1) # complete translation
    keys2, values2 = make_dictionary(path2) # not complete translation
    
    # I assume the file with more keys is the complete translation
    if len(keys1) < len(keys2):
        temp = keys1
        keys1 = keys2
        keys2 = temp
    
        temp = values1
        values1 = values2
        values2 = temp

    dict1 = dict(zip(keys1, values1))
    dict2 = dict(zip(keys2, values2)) 

    for i in dict1:
        # A translation already exists. Do nothing. 
        if i in dict2:
            continue
        # A translation does not exist. Add the original translation, which really helps. 
        dict2[i] = dict1[i]

    # They now definitely have the same keys. 
    out = []

    # KEYS1 KEYS2. IF SOMETHING IS WRONG IT'S THIS
    for i in keys1:
        if dict2[i] == "": 
            out.append(i.rstrip()) 
            continue
        out.append("%s=%s" % (i, dict2[i]))

    if has_docx:
        doc = Document()
        #doc.add_paragraph("".join(out)) # cursed
        for i in out:
            doc.add_paragraph(i.rstrip())
        doc.save(path3)
    else: 
        byproduct = open(path3, "w", encoding = "utf-8")
        byproduct.write("\n".join(out))
        byproduct.close()
        
def handle_locale(): 
    # This program will call some function on pairs of files in subfolders. See: folders, language
    
    if not os.path.isdir(path): 
        sys.exit() 
    
    # Operate on all these folders. Predictable pattern i.e. actors -> actors.properties, actors_en.properties, ...
    folders = ["actors", "items", "journal", "levels", "misc", "plants", "scenes", "ui", "windows"]
    
    # Target language. i.e. target actors.properties, actors_en.properties
    language = "es" # en
    
    for i in folders:
        destination = os.path.join(path, i)
        A = os.path.join(destination, i + ".properties")
        B = os.path.join(destination, i + "_%s.properties" % language)
        if not os.path.isdir(destination):
            print("Couldn't find this directory: " + destination)
            continue
        # I need 2 files to run the script. 
        if (not os.path.isfile(A)) or (not os.path.isfile(B)):
            print("One of these two files are missing: %s, %s" % (A, B))
            continue
        # i.e. actors_en.properties -> actors_en.docx or txt
        output = "%s_%s%s" % (i, language, suffix) 
        perform(A, B, output)
        # print("Success! Your file is located at %s\\%s." % (os.getcwd(), output))
    return "Success! Your files are located at %s." % os.getcwd()

    
